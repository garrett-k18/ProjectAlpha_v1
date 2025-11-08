"""Valuation document extraction pipeline.

This module provides an OCR-first workflow that extracts structured
valuation data from PDF, Word, Excel, or image-based BPO/Appraisal documents.
It minimises AI usage by relying on deterministic parsing first and calls an
optional Claude client only for low-confidence fields.
"""

from __future__ import annotations

import json
import logging
import mimetypes
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence

from django.utils import timezone
from core.models.valuations import Valuation

logger = logging.getLogger(__name__)

# Optional third-party dependencies -------------------------------------------------

tesseract_available = False
try:  # pragma: no cover - optional dependency
    import pytesseract
    from pytesseract import Output as TesseractOutput

    tesseract_available = True
except Exception:  # pragma: no cover - best effort optional import
    pytesseract = None
    TesseractOutput = None

pdfplumber_available = False
try:  # pragma: no cover
    import pdfplumber

    pdfplumber_available = True
except Exception:  # pragma: no cover
    pdfplumber = None

pymupdf_available = False
try:  # pragma: no cover
    import fitz  # PyMuPDF

    pymupdf_available = True
except Exception:  # pragma: no cover
    fitz = None

docx_available = False
try:  # pragma: no cover
    import docx

    docx_available = True
except Exception:  # pragma: no cover
    docx = None

pandas_available = False
try:  # pragma: no cover
    import pandas as pd

    pandas_available = True
except Exception:  # pragma: no cover
    pd = None

pil_available = False
try:  # pragma: no cover
    from PIL import Image

    pil_available = True
except Exception:  # pragma: no cover
    Image = None


# Data structures -------------------------------------------------------------------


@dataclass
class TextBlock:
    """Represents a contiguous text chunk with an optional confidence score."""

    text: str
    confidence: float = 1.0
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ExtractionOutput:
    """Normalized textual content obtained from a document."""

    full_text: str
    blocks: List[TextBlock]
    tables: List[List[Dict[str, Any]]] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    inferred_source: Optional[str] = None


@dataclass
class FieldExtractionRecord:
    """Structured result for a single extracted field."""

    model_label: str
    field: str
    value: Any
    raw_text: str
    confidence: float
    method: str
    requires_review: bool
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentExtractionResult:
    """Composite result returned by the orchestrator."""

    file_path: Path
    mime_type: str
    extracted_at: timezone.datetime
    text_output: ExtractionOutput
    fields: List[FieldExtractionRecord]
    warnings: List[str] = field(default_factory=list)
    inferred_source: Optional[str] = None


# Utility functions -----------------------------------------------------------------


def _normalise_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def _detect_mime_type(file_path: Path) -> str:
    mime, _ = mimetypes.guess_type(str(file_path))
    return mime or "application/octet-stream"


# Text extraction implementations ----------------------------------------------------


class BaseDocumentExtractor:
    """Interface for extracting textual content from a document."""

    SUPPORTED_MIME_TYPES: Sequence[str] = ()

    def supports(self, mime_type: str, file_path: Path) -> bool:
        return mime_type in self.SUPPORTED_MIME_TYPES

    def extract(self, file_path: Path) -> ExtractionOutput:
        raise NotImplementedError


class PDFDocumentExtractor(BaseDocumentExtractor):
    SUPPORTED_MIME_TYPES = ("application/pdf",)

    def extract(self, file_path: Path) -> ExtractionOutput:
        blocks: List[TextBlock] = []
        tables: List[List[Dict[str, Any]]] = []
        warnings: List[str] = []

        if pdfplumber_available:  # Prefer vector text extraction
            try:
                with pdfplumber.open(str(file_path)) as pdf:
                    for page_number, page in enumerate(pdf.pages, start=1):
                        text = page.extract_text() or ""
                        if text.strip():
                            blocks.append(
                                TextBlock(
                                    text=text,
                                    confidence=0.99,
                                    metadata={"page": page_number, "source": "pdfplumber"},
                                )
                            )
                        else:
                            warnings.append(f"No direct text on PDF page {page_number} (pdfplumber).")
                        try:
                            page_tables = page.extract_tables() or []
                            cleaned_tables = []
                            for raw in page_tables:
                                headers = raw[0]
                                for row in raw[1:]:
                                    if len(row) == len(headers):
                                        cleaned_tables.append(
                                            {headers[idx] or f"col_{idx}": row[idx] for idx in range(len(headers))}
                                        )
                            if cleaned_tables:
                                tables.append(cleaned_tables)
                        except Exception as table_err:  # pragma: no cover - best effort
                            warnings.append(f"Table extraction error on page {page_number}: {table_err}")
            except Exception as pdf_err:  # pragma: no cover
                warnings.append(f"pdfplumber failed: {pdf_err}")

        if not blocks and pymupdf_available:
            try:
                doc = fitz.open(str(file_path))
                for page_index, page in enumerate(doc, start=1):
                    text = page.get_text("text")
                    if text.strip():
                        blocks.append(
                            TextBlock(
                                text=text,
                                confidence=0.95,
                                metadata={"page": page_index, "source": "pymupdf"},
                            )
                        )
                    else:
                        warnings.append(f"No text extracted via PyMuPDF on page {page_index}.")
            except Exception as pymupdf_err:  # pragma: no cover
                warnings.append(f"PyMuPDF failed: {pymupdf_err}")

        # OCR fallback for empty or low-confidence outputs
        if tesseract_available and (not blocks or all(b.confidence < 0.5 for b in blocks)):
            if not pymupdf_available:
                warnings.append(
                    "OCR fallback requested but PyMuPDF is unavailable to render pages; skipping OCR."
                )
            else:
                try:
                    doc = fitz.open(str(file_path))
                    for page_index, page in enumerate(doc, start=1):
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scaling for OCR clarity
                        image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples) if pil_available else None
                        if image is None:
                            warnings.append("Pillow not installed; cannot convert PDF page to image for OCR.")
                            break
                        ocr_content = pytesseract.image_to_data(
                            image, output_type=TesseractOutput.DICT, lang="eng"
                        )
                        text = " ".join(
                            ocr_content["text"][i]
                            for i in range(len(ocr_content["text"]))
                            if float(ocr_content.get("conf", [0] * len(ocr_content["text"]))[i]) > 0
                        )
                        confidence_scores = [
                            float(conf)
                            for conf in ocr_content.get("conf", [])
                            if conf not in {"-1", "-0"}
                        ]
                        avg_conf = (
                            sum(confidence_scores) / len(confidence_scores) / 100.0
                            if confidence_scores
                            else 0.55
                        )
                        blocks.append(
                            TextBlock(
                                text=text,
                                confidence=avg_conf,
                                metadata={"page": page_index, "source": "ocr"},
                            )
                        )
                except Exception as ocr_err:  # pragma: no cover
                    warnings.append(f"OCR fallback failed: {ocr_err}")

        full_text = "\n".join(block.text for block in blocks)
        return ExtractionOutput(full_text=full_text, blocks=blocks, tables=tables, warnings=warnings)


class WordDocumentExtractor(BaseDocumentExtractor):
    SUPPORTED_MIME_TYPES = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )

    def extract(self, file_path: Path) -> ExtractionOutput:
        blocks: List[TextBlock] = []
        tables: List[List[Dict[str, Any]]] = []
        warnings: List[str] = []

        if not docx_available:
            warnings.append("python-docx not installed; cannot parse Word documents.")
            return ExtractionOutput(full_text="", blocks=[], warnings=warnings)

        document = docx.Document(str(file_path))
        for idx, para in enumerate(document.paragraphs, start=1):
            text = para.text.strip()
            if text:
                blocks.append(TextBlock(text=text, confidence=0.95, metadata={"paragraph": idx}))
        for table_index, table in enumerate(document.tables, start=1):
            headers = [cell.text.strip() or f"column_{i}" for i, cell in enumerate(table.rows[0].cells)]
            table_rows: List[Dict[str, Any]] = []
            for row in table.rows[1:]:
                row_dict = {}
                for i, cell in enumerate(row.cells):
                    row_dict[headers[i]] = cell.text.strip()
                table_rows.append(row_dict)
            if table_rows:
                tables.append(table_rows)
        full_text = "\n".join(block.text for block in blocks)
        return ExtractionOutput(full_text=full_text, blocks=blocks, tables=tables, warnings=warnings)


class ImageDocumentExtractor(BaseDocumentExtractor):
    SUPPORTED_MIME_TYPES = (
        "image/jpeg",
        "image/png",
        "image/tiff",
        "image/bmp",
    )

    def extract(self, file_path: Path) -> ExtractionOutput:
        warnings: List[str] = []
        if not (tesseract_available and pil_available):
            warnings.append("OCR dependencies (pytesseract + Pillow) required for image extraction.")
            return ExtractionOutput(full_text="", blocks=[], warnings=warnings)

        try:
            image = Image.open(str(file_path))
        except Exception as err:
            warnings.append(f"Failed to open image: {err}")
            return ExtractionOutput(full_text="", blocks=[], warnings=warnings)

        ocr_data = pytesseract.image_to_data(image, output_type=TesseractOutput.DICT, lang="eng")
        tokens = []
        confidences = []
        for token, conf in zip(ocr_data.get("text", []), ocr_data.get("conf", [])):
            token = token.strip()
            try:
                conf_value = float(conf)
            except (TypeError, ValueError):
                conf_value = -1.0
            if token:
                tokens.append(token)
                if conf_value >= 0:
                    confidences.append(conf_value)
        avg_conf = (sum(confidences) / len(confidences) / 100.0) if confidences else 0.6
        text = " ".join(tokens)
        block = TextBlock(text=text, confidence=avg_conf, metadata={"source": "ocr"})
        return ExtractionOutput(full_text=text, blocks=[block], warnings=warnings)


class ExcelDocumentExtractor(BaseDocumentExtractor):
    SUPPORTED_MIME_TYPES = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
        "text/csv",
    )

    def extract(self, file_path: Path) -> ExtractionOutput:
        if not pandas_available:
            warning = "pandas not installed; cannot parse spreadsheet documents."
            return ExtractionOutput(full_text="", blocks=[], warnings=[warning])

        warnings: List[str] = []
        tables: List[List[Dict[str, Any]]] = []
        blocks: List[TextBlock] = []

        suffix = file_path.suffix.lower()
        try:
            if suffix == ".csv":
                df = pd.read_csv(file_path)
                tables.append(df.fillna("").to_dict(orient="records"))
            else:
                df_dict = pd.read_excel(file_path, sheet_name=None)
                for sheet_name, df in df_dict.items():
                    table_rows = df.fillna("").to_dict(orient="records")
                    tables.append(table_rows)
                    blocks.append(
                        TextBlock(
                            text=f"Sheet {sheet_name}: {json.dumps(table_rows[:3])}",
                            confidence=0.9,
                            metadata={"sheet": sheet_name},
                        )
                    )
        except Exception as excel_err:
            warnings.append(f"Failed to parse spreadsheet: {excel_err}")

        full_text = "\n".join(block.text for block in blocks)
        return ExtractionOutput(full_text=full_text, blocks=blocks, tables=tables, warnings=warnings)


# Field extraction -------------------------------------------------------------------


@dataclass
class FieldPattern:
    model_label: str
    field: str
    pattern: re.Pattern[str]
    clean: Callable[[str], Any]
    confidence: float
    method: str = "rule"
    context_hint: str = ""


class FieldExtractionEngine:
    """Rule-based extractor that maps text snippets to model fields."""

    def __init__(self) -> None:
        self.patterns: List[FieldPattern] = [
            FieldPattern(
                model_label="etl.ValuationETL",
                field="loan_number",
                pattern=re.compile(r"(?:Loan|Order)\s*(?:Number|#)\s*[:\-]?\s*([A-Z0-9\-]+)", re.IGNORECASE),
                clean=lambda value: value.strip().upper(),
                confidence=0.99,
                context_hint="Loan Number",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="inspection_date",
                pattern=re.compile(
                    r"(?:Inspection|Effective|Report)\s*(?:Date|Dt\.)\s*[:\-]?\s*(\d{1,2}[\-/]\d{1,2}[\-/]\d{2,4})",
                    re.IGNORECASE,
                ),
                clean=lambda value: value.strip(),
                confidence=0.95,
                context_hint="Inspection Date",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="property_address",
                pattern=re.compile(
                    r"(?:Subject\s+Property\s+)?Address\s*[:\-]?\s*(.+)",
                    re.IGNORECASE,
                ),
                clean=lambda value: _normalise_text(value),
                confidence=0.9,
                context_hint="Subject Property Address",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="city",
                pattern=re.compile(r"City\s*[:\-]?\s*([A-Za-z\s]+)", re.IGNORECASE),
                clean=lambda value: _normalise_text(value),
                confidence=0.85,
                context_hint="City",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="state",
                pattern=re.compile(r"State\s*[:\-]?\s*([A-Z]{2})\b"),
                clean=lambda value: value.strip().upper(),
                confidence=0.9,
                context_hint="State",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="zip_code",
                pattern=re.compile(r"Zip\s*(?:Code)?\s*[:\-]?\s*(\d{5}(?:-\d{4})?)", re.IGNORECASE),
                clean=lambda value: value.strip(),
                confidence=0.92,
                context_hint="Zip",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="valuation_type",
                pattern=re.compile(r"Valuation\s*Type\s*[:\-]?\s*(\w+)", re.IGNORECASE),
                clean=lambda value: value.strip().upper(),
                confidence=0.8,
                context_hint="Valuation Type",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="as_is_value",
                pattern=re.compile(
                    r"As\s*[- ]?Is\s*(?:Value|Estimate)\s*[:\-]?\s*\$?([0-9,]+(?:\.\d{2})?)",
                    re.IGNORECASE,
                ),
                clean=lambda value: float(value.replace(",", "")),
                confidence=0.9,
                context_hint="As-Is Value",
            ),
            FieldPattern(
                model_label="etl.ValuationETL",
                field="as_repaired_value",
                pattern=re.compile(
                    r"As\s*[- ]?Repaired\s*(?:Value|Estimate)\s*[:\-]?\s*\$?([0-9,]+(?:\.\d{2})?)",
                    re.IGNORECASE,
                ),
                clean=lambda value: float(value.replace(",", "")),
                confidence=0.85,
                context_hint="As-Repaired Value",
            ),
        ]
        self.comparable_column_map = {
            "address": "address",
            "property address": "address",
            "subject address": "address",
            "comp address": "address",
            "city": "city",
            "state": "state",
            "st": "state",
            "zip": "zip_code",
            "zip code": "zip_code",
            "postal code": "zip_code",
            "sale price": "sale_price",
            "sold price": "sale_price",
            "closing price": "sale_price",
            "list price": "current_list_price",
            "current list price": "current_list_price",
            "original list price": "original_list_price",
            "dom": "days_on_market",
            "days on market": "days_on_market",
            "cdom": "cumulative_days_on_market",
            "distance": "proximity_miles",
            "proximity": "proximity_to_subject",
            "sqft": "living_area",
            "gla": "living_area",
            "living area": "living_area",
            "bed": "bedrooms",
            "beds": "bedrooms",
            "bedrooms": "bedrooms",
            "bath": "bathrooms",
            "baths": "bathrooms",
            "bathrooms": "bathrooms",
            "year built": "year_built",
            "comp type": "comp_type",
            "type": "comp_type",
            "sales type": "sales_type",
            "financing": "financing_type",
            "financing type": "financing_type",
        }
        self.repair_column_map = {
            "repair": "description",
            "repair item": "description",
            "repair description": "description",
            "item": "description",
            "scope": "description",
            "category": "category",
            "repair category": "category",
            "type": "repair_type",
            "location": "repair_type",
            "severity": "severity",
            "priority": "priority",
            "required": "is_required",
            "mandatory": "is_required",
            "cost": "estimated_cost",
            "estimate": "estimated_cost",
            "amount": "estimated_cost",
        }
        self.comparable_column_map = {
            "address": "address",
            "property address": "address",
            "subject address": "address",
            "comp address": "address",
            "city": "city",
            "state": "state",
            "st": "state",
            "zip": "zip_code",
            "zip code": "zip_code",
            "postal code": "zip_code",
            "sale price": "sale_price",
            "sold price": "sale_price",
            "closing price": "sale_price",
            "list price": "current_list_price",
            "current list price": "current_list_price",
            "original list price": "original_list_price",
            "dom": "days_on_market",
            "days on market": "days_on_market",
            "cdom": "cumulative_days_on_market",
            "distance": "proximity_miles",
            "proximity": "proximity_to_subject",
            "sqft": "living_area",
            "gla": "living_area",
            "living area": "living_area",
            "bed": "bedrooms",
            "beds": "bedrooms",
            "bedrooms": "bedrooms",
            "bath": "bathrooms",
            "baths": "bathrooms",
            "bathrooms": "bathrooms",
            "year built": "year_built",
            "comp type": "comp_type",
            "type": "comp_type",
            "sales type": "sales_type",
            "financing": "financing_type",
            "financing type": "financing_type",
        }
        self.repair_column_map = {
            "repair": "description",
            "repair item": "description",
            "repair description": "description",
            "item": "description",
            "scope": "description",
            "category": "category",
            "repair category": "category",
            "type": "repair_type",
            "location": "repair_type",
            "severity": "severity",
            "priority": "priority",
            "required": "is_required",
            "mandatory": "is_required",
            "cost": "estimated_cost",
            "estimate": "estimated_cost",
            "amount": "estimated_cost",
        }

    def extract_fields(self, text_output: ExtractionOutput) -> List[FieldExtractionRecord]:
        text = text_output.full_text
        records: List[FieldExtractionRecord] = []
        seen_keys = set()

        for pattern in self.patterns:
            match = pattern.pattern.search(text)
            if not match:
                continue
            value = pattern.clean(match.group(1))
            key = (pattern.model_label, pattern.field)
            if key in seen_keys:
                continue
            record = FieldExtractionRecord(
                model_label=pattern.model_label,
                field=pattern.field,
                value=value,
                raw_text=_normalise_text(match.group(0)),
                confidence=pattern.confidence,
                method=pattern.method,
                requires_review=pattern.confidence < 0.8,
                metadata={"context_hint": pattern.context_hint},
            )
            records.append(record)
            seen_keys.add(key)

        # Attempt to parse comparables from tables when supplied
        for table_index, table_rows in enumerate(text_output.tables):
            if not table_rows:
                continue
            normalized_headers = self._normalise_headers(table_rows[0])
            if self._looks_like_comparable_table(normalized_headers):
                records.extend(
                    self._extract_comparable_records(table_rows, table_index)
                )
                continue
            if self._looks_like_repair_table(normalized_headers):
                records.extend(
                    self._extract_repair_records(table_rows, table_index)
                )
                continue

        return records

    @staticmethod
    def _normalise_headers(sample_row: Dict[str, Any]) -> Dict[str, str]:
        return {
            (key or "").strip().lower(): key for key in sample_row.keys() if key is not None
        }

    def _looks_like_comparable_table(self, headers: Dict[str, str]) -> bool:
        header_keys = set(headers.keys())
        return "address" in header_keys and (
            "sale price" in header_keys
            or "sold price" in header_keys
            or "list price" in header_keys
            or "closing price" in header_keys
        )

    def _looks_like_repair_table(self, headers: Dict[str, str]) -> bool:
        header_keys = set(headers.keys())
        return (
            "repair" in header_keys
            or "repair item" in header_keys
            or "repair description" in header_keys
        ) and ("cost" in header_keys or "estimate" in header_keys or "amount" in header_keys)

    def _extract_comparable_records(
        self,
        table_rows: Sequence[Dict[str, Any]],
        table_index: int,
    ) -> List[FieldExtractionRecord]:
        comparable_records: List[FieldExtractionRecord] = []
        for row_index, row in enumerate(table_rows):
            comp_data: Dict[str, Any] = {}
            for key, value in row.items():
                if key is None:
                    continue
                normalized_key = key.strip().lower()
                field_name = self.comparable_column_map.get(normalized_key)
                if not field_name:
                    continue
                if isinstance(value, str):
                    value = value.strip()
                comp_data[field_name] = value
            if not comp_data.get("address"):
                continue
            comp_data.setdefault("comp_number", row_index + 1)
            confidence = 0.8 if comp_data.get("sale_price") else 0.6
            comparable_records.append(
                FieldExtractionRecord(
                    model_label="etl.ComparablesETL",
                    field=f"comparable_{row_index + 1}",
                    value=comp_data,
                    raw_text=json.dumps(row),
                    confidence=confidence,
                    method="table",
                    requires_review=confidence < 0.85,
                    metadata={
                        "table_index": table_index,
                        "row_index": row_index,
                    },
                )
            )
        return comparable_records

    def _extract_repair_records(
        self,
        table_rows: Sequence[Dict[str, Any]],
        table_index: int,
    ) -> List[FieldExtractionRecord]:
        repair_records: List[FieldExtractionRecord] = []
        for row_index, row in enumerate(table_rows):
            repair_data: Dict[str, Any] = {}
            for key, value in row.items():
                if key is None:
                    continue
                normalized_key = key.strip().lower()
                field_name = self.repair_column_map.get(normalized_key)
                if not field_name:
                    continue
                if isinstance(value, str):
                    value = value.strip()
                repair_data[field_name] = value
            if not repair_data:
                continue
            repair_data.setdefault("repair_number", row_index + 1)
            confidence = 0.7 if repair_data.get("estimated_cost") else 0.55
            repair_records.append(
                FieldExtractionRecord(
                    model_label="etl.RepairItem",
                    field=f"repair_{row_index + 1}",
                    value=repair_data,
                    raw_text=json.dumps(row),
                    confidence=confidence,
                    method="table",
                    requires_review=True,
                    metadata={
                        "table_index": table_index,
                        "row_index": row_index,
                    },
                )
            )
        return repair_records


# Claude augmentation ---------------------------------------------------------------


class ClaudeFieldAugmenter:
    """Optional integration with Claude for low-confidence fields.

    The caller must provide a callable that accepts a prompt (str) and returns a
    dictionary with `value` and `confidence` keys (additional metadata optional).
    """

    def __init__(
        self,
        client: Optional[Callable[[str, Dict[str, Any]], Dict[str, Any]]] = None,
        model_name: str = "claude-3-haiku-20240307",
    ) -> None:
        self.client = client
        self.model_name = model_name

    def augment(  # pragma: no cover - depends on external API
        self,
        text_output: ExtractionOutput,
        pending_records: Sequence[FieldExtractionRecord],
    ) -> List[FieldExtractionRecord]:
        if not self.client or not pending_records:
            return []

        augmented: List[FieldExtractionRecord] = []
        for record in pending_records:
            prompt = self._build_prompt(text_output, record)
            try:
                result = self.client(prompt, {"model": self.model_name, "field": record.field})
            except Exception as err:
                logger.warning("Claude augmentation failed for %s: %s", record.field, err)
                continue
            if not result:
                continue
            value = result.get("value")
            confidence = float(result.get("confidence", 0.6))
            if value in (None, ""):
                continue
            augmented.append(
                FieldExtractionRecord(
                    model_label=record.model_label,
                    field=record.field,
                    value=value,
                    raw_text=result.get("raw_text", record.raw_text),
                    confidence=confidence,
                    method="ai",
                    requires_review=confidence < 0.9,
                    metadata={"augmenter": "claude", **result.get("metadata", {})},
                )
            )
        return augmented

    @staticmethod
    def _build_prompt(text_output: ExtractionOutput, record: FieldExtractionRecord) -> str:
        context_snippets = [block.text for block in text_output.blocks[:5]]
        context = "\n\n".join(context_snippets)
        return (
            "You are extracting structured data from an appraisal/BPO document.\n"
            f"Field: {record.field}\n"
            "Return JSON with keys `value` and `confidence` (0-1)."
            "If unsure, respond with an empty value.\n"
            "Document excerpt:\n"
            f"{context}\n"
        )


# Orchestrator ----------------------------------------------------------------------


class DocumentExtractionService:
    """Coordinates document parsing, field extraction, and AI fallback."""

    def __init__(
        self,
        text_extractors: Optional[Sequence[BaseDocumentExtractor]] = None,
        field_engine: Optional[FieldExtractionEngine] = None,
        claude_augmenter: Optional[ClaudeFieldAugmenter] = None,
        high_confidence_threshold: float = 0.98,
        ai_threshold: float = 0.85,
        review_threshold: float = 0.75,
    ) -> None:
        self.text_extractors: List[BaseDocumentExtractor] = list(
            text_extractors
            or [
                PDFDocumentExtractor(),
                WordDocumentExtractor(),
                ExcelDocumentExtractor(),
                ImageDocumentExtractor(),
            ]
        )
        self.field_engine = field_engine or FieldExtractionEngine()
        self.claude_augmenter = claude_augmenter or ClaudeFieldAugmenter()
        self.high_confidence_threshold = high_confidence_threshold
        self.ai_threshold = ai_threshold
        self.review_threshold = review_threshold

    # Public API ------------------------------------------------------------------

    def process(self, file_path: Path) -> DocumentExtractionResult:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(file_path)

        mime_type = _detect_mime_type(file_path)
        extractor = self._select_extractor(mime_type, file_path)
        if extractor is None:
            raise ValueError(f"Unsupported file type: {mime_type}")

        logger.info("Extracting text from %s using %s", file_path, extractor.__class__.__name__)
        text_output = extractor.extract(file_path)
        warnings = list(text_output.warnings)

        field_records = self.field_engine.extract_fields(text_output)
        logger.debug("Rule-based extraction produced %d fields", len(field_records))

        inferred_source = self._infer_source(text_output)
        text_output.inferred_source = inferred_source
        if inferred_source:
            logger.debug("Inferred valuation source %s from document text", inferred_source)

        low_confidence_records = [
            record for record in field_records if record.confidence < self.ai_threshold
        ]
        if low_confidence_records:
            ai_records = self.claude_augmenter.augment(text_output, low_confidence_records)
            field_records = self._merge_records(field_records, ai_records)

        for record in field_records:
            # Flag for manual review when confidence is low and no AI augmentation improved it
            if record.confidence < self.review_threshold:
                record.requires_review = True
            else:
                record.requires_review = record.requires_review or False

            if record.confidence >= self.high_confidence_threshold:
                record.metadata.setdefault("finalised", True)

        return DocumentExtractionResult(
            file_path=file_path,
            mime_type=mime_type,
            extracted_at=timezone.now(),
            text_output=text_output,
            fields=field_records,
            warnings=warnings,
            inferred_source=inferred_source,
        )

    # Internal helpers -------------------------------------------------------------

    def _select_extractor(
        self, mime_type: str, file_path: Path
    ) -> Optional[BaseDocumentExtractor]:
        for extractor in self.text_extractors:
            if extractor.supports(mime_type, file_path):
                return extractor
        # Handle ambiguous types (e.g., .pdf with octet-stream mime)
        suffix = file_path.suffix.lower()
        for extractor in self.text_extractors:
            if suffix == ".pdf" and isinstance(extractor, PDFDocumentExtractor):
                return extractor
            if suffix in {".doc", ".docx"} and isinstance(extractor, WordDocumentExtractor):
                return extractor
            if suffix in {".xls", ".xlsx", ".csv"} and isinstance(extractor, ExcelDocumentExtractor):
                return extractor
            if suffix in {".jpg", ".jpeg", ".png", ".tif", ".tiff"} and isinstance(
                extractor, ImageDocumentExtractor
            ):
                return extractor
        return None

    def _infer_source(self, text_output: ExtractionOutput) -> Optional[str]:
        """Attempt to infer Valuation.Source choice from document text."""

        text = (text_output.full_text or "").lower()
        if not text:
            return None

        if "broker price opinion" in text:
            if "interior" in text:
                return Valuation.Source.BPO_INTERIOR
            if "exterior" in text:
                return Valuation.Source.BPO_EXTERIOR
            return Valuation.Source.BROKER

        if any(keyword in text for keyword in ["appraisal report", "summary appraisal", "uniform residential appraisal"]):
            return Valuation.Source.APPRAISAL

        if "broker opinion" in text:
            return Valuation.Source.BROKER

        if "desktop valuation" in text:
            return Valuation.Source.DESKTOP

        if "internal valuation" in text:
            return Valuation.Source.INTERNAL

        return None

    @staticmethod
    def _merge_records(
        base_records: Sequence[FieldExtractionRecord],
        overrides: Sequence[FieldExtractionRecord],
    ) -> List[FieldExtractionRecord]:
        merged = {(record.model_label, record.field): record for record in base_records}
        for override in overrides:
            key = (override.model_label, override.field)
            if key not in merged or override.confidence >= merged[key].confidence:
                merged[key] = override
        return list(merged.values())


__all__ = [
    "DocumentExtractionService",
    "FieldExtractionEngine",
    "FieldExtractionRecord",
    "DocumentExtractionResult",
    "ClaudeFieldAugmenter",
    "ExtractionOutput",
    "TextBlock",
]
